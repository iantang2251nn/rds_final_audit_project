"""Generate report_template.docx — section skeleton mirroring report.tex."""
from pathlib import Path
from docx import Document

OUT = Path(__file__).parent / "report_template.docx"

STRUCTURE = [
    (1, "Abstract"),
    (1, "1. Introduction"),
    (1, "2. Background"),
    (2, "2.1 The Civil Comments dataset and Jigsaw competition"),
    (2, "2.2 The 3rd-place solution under audit"),
    (2, "2.3 Normative framing: fairness as the spine of the audit"),
    (1, "3. Input Data Profiling"),
    (2, "3.1 Dataset composition"),
    (2, "3.2 Identity-subgroup coverage and base rates"),
    (2, "3.3 Critique of the toxicity-score definition"),
    (1, "4. The ADS Implementation"),
    (2, "4.1 Recreation scope and deviations from the original"),
    (2, "4.2 Architecture: BERT-large + LSTM-GRU multi-task ensemble"),
    (2, "4.3 Training and Optuna-blended weights"),
    (1, "5. Stakeholder Analysis"),
    (1, "6. Audit Outcomes"),
    (2, "6.1 Pillar 1 — Subgroup fairness analysis"),
    (3, "6.1.1 Subgroup, BPSN, and BNSP AUCs"),
    (3, "6.1.2 Direction of error: over- vs. under-flagging"),
    (2, "6.2 Pillar 2 — Counterfactual identity perturbation"),
    (3, "6.2.1 Causal DAG and the anti-causal structure of toxicity classification"),
    (3, "6.2.2 Conceptualizing the do-intervention on identity tokens"),
    (3, "6.2.3 Empirical test of the substitution-validity assumption"),
    (3, "6.2.4 Template-level results: heatmaps, marginal effects, pairwise gaps"),
    (3, "6.2.5 Model-comparison: BERT vs. LSTM vs. ensemble"),
    (2, "6.3 Pillar 3 — Module 1 fairness metrics"),
    (3, "6.3.1 Threshold choice s_HR"),
    (3, "6.3.2 Per-subgroup PPV, TPR, FPR, FNR, and selection rate"),
    (3, "6.3.3 Numerical demonstration of the impossibility result"),
    (2, "6.4 Pillar 4 — Critique of the competition metric"),
    (3, "6.4.1 Power-mean exponent p = -5 as an implicit Rawlsian choice"),
    (3, "6.4.2 The 25/75 weighting between overall AUC and bias metrics"),
    (3, "6.4.3 Identity column inclusion and exclusion decisions"),
    (3, "6.4.4 Meta-point: metric design as upstream value-laden choice"),
    (1, "7. Reflection and Recommendations"),
    (2, "7.1 Was the data appropriate?"),
    (2, "7.2 Is the implementation robust, accurate, and fair?"),
    (2, "7.3 Would we deploy this in the public sector or industry?"),
    (2, "7.4 Recommendations"),
    (1, "8. Limitations"),
    (2, "8.1 Recreation scope"),
    (2, "8.2 Toxicity-label measurement"),
    (2, "8.3 Counterfactual methodology"),
    (1, "References"),
    (1, "Appendix A. Counterfactual Fairness Framework"),
    (1, "Appendix B. Module 1 Fairness Metrics"),
]


def main() -> None:
    doc = Document()

    title = doc.add_heading(
        "Auditing the 3rd-Place Solution to the Jigsaw Unintended "
        "Bias in Toxicity Classification Competition",
        level=0,
    )
    title.alignment = 1  # center

    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("Ian Tang and Nick Chen\n").bold = True
    p.add_run("DS-UA 202, Spring 2026 — Group 21")

    doc.add_paragraph()

    for level, heading in STRUCTURE:
        doc.add_heading(heading, level=level)
        doc.add_paragraph()  # placeholder paragraph under each heading

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
